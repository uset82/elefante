import os, re
from docx import Document
from docx.shared import Inches

ROOT = r"d:\HVL2025\ADA525\uno"
REPORT_DIR = os.path.join(ROOT, "report")
MD_PATH = os.path.join(REPORT_DIR, "reportchatgpt.md")
OUT_PATH = os.path.join(REPORT_DIR, "reportchatgpt.docx")

# helper: resolve relative image paths
_DEF_WIDTH_IN = 6.5
_img_counter = 0

def resolve_path(rel):
    if rel.startswith('./') or rel.startswith('.\\'):
        rel = rel[2:]
    return os.path.normpath(os.path.join(REPORT_DIR, rel))


def add_image(doc, path):
    global _img_counter
    _img_counter += 1
    ext = os.path.splitext(path)[1].lower()
    if ext == '.svg':
        # convert to PNG with cairosvg if available, else skip
        try:
            from cairosvg import svg2png
            png_path = os.path.join(REPORT_DIR, f"_md_img_{_img_counter}.png")
            svg2png(url=path, write_to=png_path, output_width=1600)
            doc.add_picture(png_path, width=Inches(_DEF_WIDTH_IN))
        except Exception:
            # fallback: ignore or insert a paragraph placeholder
            doc.add_paragraph(f"[Image placeholder: {os.path.basename(path)}]")
    else:
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(_DEF_WIDTH_IN))
        else:
            doc.add_paragraph(f"[Missing image: {os.path.basename(path)}]")


def convert_md_to_docx(md_path, out_path):
    doc = Document()
    # Add a visible revision timestamp at the top so you can confirm updates
    from datetime import datetime
    doc.add_paragraph(f"Revision: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    if not os.path.exists(md_path):
        raise FileNotFoundError(md_path)

    with open(md_path, 'r', encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\n')
            # image: ![alt](path)
            m_img = re.match(r"!\[[^\]]*\]\(([^\)]+)\)", line)
            if m_img:
                img_rel = m_img.group(1)
                add_image(doc, resolve_path(img_rel))
                continue
            # headings
            m_h = re.match(r"^(#{1,6})\s+(.*)$", line)
            if m_h:
                level = len(m_h.group(1))
                text = m_h.group(2)
                # map H1..H6 to docx headings (cap at 3)
                level = min(level, 3)
                doc.add_heading(text, level=level)
                continue
            # bullet list
            if line.strip().startswith('- '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
                continue
            # numbered list
            m_num = re.match(r"^\d+\.\s+(.*)$", line.strip())
            if m_num:
                doc.add_paragraph(m_num.group(1), style='List Number')
                continue
            # blank line -> spacing
            if line.strip() == '':
                doc.add_paragraph('')
                continue
            # default paragraph
            doc.add_paragraph(line)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    os.makedirs(REPORT_DIR, exist_ok=True)
    convert_md_to_docx(MD_PATH, OUT_PATH)