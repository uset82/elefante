from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

ROOT = r"d:\HVL2025\ADA525\uno"
REPORT_DIR = os.path.join(ROOT, "report")
MOCKUP_DIR = os.path.join(ROOT, "mockup")

# Input assets
SCREENSHOT = os.path.join(REPORT_DIR, "Screenshot 2025-10-24 165700.png")
UI_SVG = os.path.join(MOCKUP_DIR, "ui-components.svg")
SYS_SVG_FALLBACK_MOCKUP = os.path.join(MOCKUP_DIR, "system-diagram.svg")
SYS_SVG_REPORT = os.path.join(REPORT_DIR, "system-diagram.svg")

# Output PNGs
UI_PNG = os.path.join(REPORT_DIR, "ui-components.png")
SYS_PNG = os.path.join(REPORT_DIR, "system-diagram.png")


def ensure_pngs():
    """Convert SVGs to PNGs when possible; otherwise create simple placeholders."""
    # UI components
    try:
        from cairosvg import svg2png
        if os.path.exists(UI_SVG):
            svg2png(url=UI_SVG, write_to=UI_PNG, output_width=1600)
        else:
            raise FileNotFoundError(UI_SVG)
    except Exception:
        from PIL import Image, ImageDraw
        img = Image.new("RGB", (1600, 1000), "#0f0816")
        d = ImageDraw.Draw(img)
        d.text((60, 60), "UI Components (placeholder)", fill="#e6dbff")
        img.save(UI_PNG)

    # System diagram — prefer report/system-diagram.svg if present
    try:
        from cairosvg import svg2png
        if os.path.exists(SYS_SVG_REPORT):
            svg2png(url=SYS_SVG_REPORT, write_to=SYS_PNG, output_width=1600)
        elif os.path.exists(SYS_SVG_FALLBACK_MOCKUP):
            svg2png(url=SYS_SVG_FALLBACK_MOCKUP, write_to=SYS_PNG, output_width=1600)
        else:
            raise FileNotFoundError("system-diagram.svg not found")
    except Exception:
        from PIL import Image, ImageDraw
        img = Image.new("RGB", (1600, 1000), "#0f0816")
        d = ImageDraw.Draw(img)
        d.text((60, 60), "System Diagram (placeholder)", fill="#e6dbff")
        img.save(SYS_PNG)


def build_docx():
    doc = Document()

    # Title
    doc.add_heading("Elephant Watering System — Iteration Report", 0)
    doc.add_paragraph("Date: Wednesday 29.10")

    # Figure 1 — Screenshot of real signals (if available)
    if os.path.exists(SCREENSHOT):
        p0 = doc.add_paragraph(); r0 = p0.add_run(); r0.add_picture(SCREENSHOT, width=Inches(6.5)); p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 1. Live plotter screenshot — shows Temperature, Humidity, Soil ADC/% and Motor/PWM charts with current values.")

    # Figure 2 — UI Components (dark glassy)
    p1 = doc.add_paragraph(); r1 = p1.add_run(); r1.add_picture(UI_PNG, width=Inches(6.5)); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 2. Print‑ready UI components — dark glassy card layout mirroring the live plotter for paper mockup.")

    # Figure 3 — System Diagram
    p2 = doc.add_paragraph(); r2 = p2.add_run(); r2.add_picture(SYS_PNG, width=Inches(6.5)); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 3. Hardware wiring overview — DHT11 on D2, soil sensor AOUT on A0, relay on D8 controlling an externally powered pump.")

    # Teacher rubric sections (≤ 750 words)
    doc.add_heading("Your Key Design Decisions", level=1)
    for item in [
        "Simple, safe wiring: Uno switches a relay on D8; pump gets external power.",
        "Watering only on stable dry readings: hysteresis + stability timers + minimum ON/OFF + MAX_ON_TIME safety cutoff.",
        "Dual soil signal representation: raw ADC for calibration and derived moisture % for quick understanding (DRY≈1023, WET≈300 mapping).",
        "Keep runtime UI focused on signals; place flowchart and mockups in separate folders for printing and explanation.",
        "Use SVG/PNG assets so A4 prints are crisp and easy to update."
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Your Current Implementation Progress", level=1)
    for item in [
        "Firmware samples DHT11 (~2s) and reads A0 with low‑pass filter (soilAvg = 0.9*soilAvg + 0.1*soilValue).",
        "Serial output matches plotter parser: Temp | Pot | Servo | Motor | Humidity | ON/OFF.",
        "Plotter renders Temperature, Humidity, Soil ADC, Soil %, Servo, and Motor/PWM with current values and pump state.",
        "Mermaid flowchart documents sampling, smoothing, dry/wet detection, pump ON/OFF, and safety cutoff.",
        "Print assets ready: plotter-wireframe (mockup), UI components, and system diagram."
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("The Final Outcome of This Iteration", level=1)
    doc.add_paragraph(
        "A cohesive physical mockup and a working plotter aligned with firmware behavior. The mockup mirrors the live UI and the flowchart + wiring diagram make it easy to explain pump decisions and safety in the studio.")

    doc.add_heading("New Elements and Improvements", level=1)
    for item in [
        "Moisture percentage view (derived from ADC) improves readability for non‑technical reviewers.",
        "Dark glassy theme increases legibility and maintains visual continuity between screen and paper.",
        "Stability safeguards reduce relay chatter and provide predictable watering behavior.",
        "Documentation assets (component sheet, flowchart, diagram) streamline discussion and feedback during presentation.",
        "Planned: add a rotating base in the next version to support multiple plants."  # from your note
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Save
    out_path = os.path.join(REPORT_DIR, "studio-report.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    os.makedirs(REPORT_DIR, exist_ok=True)
    ensure_pngs()
    build_docx()